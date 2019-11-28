from docx import Document
from docx.shared import Inches,Mm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

document = Document()
sections = document.sections
section = sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)

section.top_margin=Inches(0.7)
section.bottom_margin=Inches(0.7)
section.right_margin=Inches(0.9)
section.left_margin=Inches(0.7)
def list_number(doc, par, prev=None, level=None, num=True):
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level

 
style = document.styles['Normal']
font = style.font
font.name = 'Arial'

heading="National Electric Power Regulatory Authority\n(NEPRA)\nIslamabad–Pakistan"
paragraph = document.add_paragraph()
 
 
 
p=paragraph
p.add_run(heading).bold = True
p.add_run("\nGENERATION LICENCE\nNo. DGL//2019").bold=True

p.left_indent = Inches(0.5)
 
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
 

p0 = document.add_paragraph(style = 'List Number',text = '\tThe Authority hereby grants  Generation Licence to Mr. Syed  Ehsan Ullah Waqas, for 26.40 KW photovoltaic solar based distributed generation facility, having consumer reference number  24-11274-9104406, located at The Punjab School Mosque, Admin Block, Sector- C, Johar Town Lahore under the National Electric Power Regulatory Authority (Alternative & Renewable Energy) Distributed Generation and Net Metering Regulations, 2015 (the “A&RE Regulations”) for a period of seven (07) years. This Licence is valid up to ___ day of               2026. ')
p0.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
list_number(document, p0, level=0, num=True)
p0.style.font.size=Pt(13)
 
 
run = document.add_paragraph().add_run()
font = run.font
font.size = Pt(13)
run.bold=True
font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
#run.text="hehhe"
run.bold=False

p2 = document.add_paragraph().add_run()
#p2.text="AOA"
 
p1 = document.add_paragraph(style = 'List Number', text = "\tThe Licensee shall abide by the provisions under the A&RE Regulations during the currency of the Generation Licence.")
list_number(document, p1,p0, level=0  )
p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
 
three="\tThe technical parameters of Net Metering arrangements are shown hereunder:-"
p2 = document.add_paragraph(style = 'List Number', text = three)
list_number(document, p2,p1,  level=0)

p3 = document.add_paragraph(style = 'List Number 3', text = "\tPrimary Energy Source:\t\tSolar")
list_number(document, p3  ,level=5,num=False)
 

p4 = document.add_paragraph(style = 'List Number 3', text = "\tSize of Distributed Generation Facility:")
list_number(document,p4,p3  , level=0)
p5 = document.add_paragraph(style = 'List Number 3', text = "\tGenerator/Inverter Information:	  ")
list_number(document,p5 ,p4 , level=0)
#(iii).	Generator/Inverter Information: 
p6 = document.add_paragraph(style = 'Normal', text = "\t\tManufacture:	  ")
#list_number(document,p3,  level=0)
#Manufacture:


p7 = document.add_paragraph(style = 'Normal', text = "\t\tModel No:	  ")
#list_number(document,p5,p4,  level=0)
p8 = document.add_paragraph(style = 'List Number 3', text = "\tGeneration Type:	  ")
list_number(document,p8,p5,  level=0)
p9 = document.add_paragraph(style = 'List Number', text = "\tThis Licence may be renewed subject to the A&RE Regulations ")
list_number(document,p9,p2,  level=0)
#Model No.
#Generation Type:
#4.	This Licence may be renewed subject to the A&RE Regulations
document.save('demo2.docx')